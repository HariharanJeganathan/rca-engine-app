from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import pandas as pd
import tempfile
import os
import asyncio
import json
import re
from datetime import datetime
from typing import Optional

# For PDF and DOCX
import fitz  # PyMuPDF
from docx import Document  # python-docx

from memory_store import RCAStore
from rca_engine import RCAAgent
from config import MAX_FILE_SIZE_MB, MAX_EXCEL_ROWS

app = FastAPI(title="Enterprise RCA Engine", version="5.1")

os.makedirs("static", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

agent = RCAAgent()
store = RCAStore()
active_sessions = {}

# ============================================================================
# UTILITIES
# ============================================================================

@app.get("/")
async def root():
    return FileResponse("static/index.html")

def process_excel(file_bytes: bytes) -> str:
    """Extract text from Excel"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    
    try:
        try:
            df = pd.read_excel(path, dtype=str, engine='openpyxl')
        except:
            df = pd.read_excel(path, dtype=str, engine='xlrd')
    except Exception as e:
        raise HTTPException(400, f"Excel read failed: {str(e)}")
    finally:
        if os.path.exists(path):
            os.remove(path)
    
    df = df.head(MAX_EXCEL_ROWS)
    parts = []
    
    for col in df.columns:
        if any(word in str(col).lower() for word in ['description', 'summary', 'detail', 'notes', 'comment', 'narrative']):
            for val in df[col]:
                if pd.notna(val) and len(str(val)) > 20:
                    parts.append(str(val))
    
    if not parts:
        for _, row in df.iterrows():
            row_text = " | ".join([f"{col}: {row[col]}" for col in df.columns 
                                   if pd.notna(row[col]) and len(str(row[col])) > 3])
            if row_text:
                parts.append(row_text)
    
    return "\n\n".join(parts) if parts else str(df.to_string())

def process_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(file_bytes)
        path = tmp.name
    
    try:
        text = ""
        with fitz.open(path) as pdf:
            for page in pdf:
                text += page.get_text()
        return text.strip()
    except Exception as e:
        raise HTTPException(400, f"PDF read failed: {str(e)}")
    finally:
        if os.path.exists(path):
            os.remove(path)

def process_docx(file_bytes: bytes) -> str:
    """Extract text from Word document (.docx)"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        path = tmp.name

    try:
        doc = Document(path)
        full_text = []

        # 1. Read paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # 2. Read tables (VERY IMPORTANT – many MIR docs use tables)
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    full_text.append(" | ".join(row_data))

        combined_text = "\n".join(full_text).strip()

        if not combined_text:
            raise Exception("No readable text found in DOCX")

        return combined_text

    except Exception as e:
        raise HTTPException(400, f"DOCX read failed: {str(e)}")

    finally:
        if os.path.exists(path):
            os.remove(path)


def process_file(file_bytes: bytes, filename: str) -> str:
    """Auto-detect file type: xlsx, xls, pdf, csv, docx"""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return process_pdf(file_bytes)
    elif filename_lower.endswith(('.xlsx', '.xls', '.csv')):
        return process_excel(file_bytes)
    elif filename_lower.endswith('.docx'):
        return process_docx(file_bytes)
    else:
        try:
            return file_bytes.decode('utf-8')
        except:
            raise HTTPException(400, f"Unsupported file format: {filename}. Allowed: .xlsx, .xls, .pdf, .csv, .docx")

# ============================================================================
# CREATE RCA (DRAFT) - With DB Lookup for existing incidents
# ============================================================================

@app.post("/create-draft")
async def create_draft(
    incident_type: str = Form(...),
    text_input: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    incident_id_manual: Optional[str] = Form(None),
    confirm_save: bool = Form(False)
):
    """
    Create draft RCA. 
    If confirm_save=False: Generate preview only (status=DRAFT)
    If confirm_save=True: Save to DB as "CONFIRMED" and return same output next time
    Only CONFIRMED RCAs are retrieved from cache on future runs.
    """
    print(f"\n{'='*60}")
    print(f"🆕 CREATE DRAFT: {incident_type} | Confirm: {confirm_save}")
    
    # Get content from file or text
    content = ""
    if file:
        print(f"📁 File: {file.filename}")
        file_bytes = await file.read()
        if len(file_bytes) > MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(400, f"File too large (max {MAX_FILE_SIZE_MB}MB)")
        
        loop = asyncio.get_event_loop()
        content = await loop.run_in_executor(
            None, 
            lambda: process_file(file_bytes, file.filename)
        )
    elif text_input:
        print(f"📝 Manual text: {len(text_input)} chars")
        content = text_input
    else:
        raise HTTPException(400, "Provide either file or text_input")
    
    if not content or len(content.strip()) < 50:
        raise HTTPException(400, "Insufficient incident data")
    
    # Extract ID from content
    inc_match = re.search(r'(INC\d{6,12})', content, re.IGNORECASE)
    incident_id = incident_id_manual or (inc_match.group(1).upper() if inc_match else "UNKNOWN")
    
    print(f"🆔 Extracted ID: {incident_id}")
    
    # ============================================================
    # CHECK IF CONFIRMED RCA ALREADY EXISTS (Only CONFIRMED, not DRAFT)
    # ============================================================
    if not confirm_save:  # Only check cache if not explicitly confirming
        existing_rca = store.get_by_incident_id(incident_id)
        # CRITICAL FIX: Check if exists AND status is CONFIRMED
        if existing_rca and existing_rca["status"] == "CONFIRMED":
            print(f"✅ Found CONFIRMED RCA for {incident_id}!")
            print(f"   Returning stored version: {existing_rca['id']}")
            
            # FIXED: Use ["key"] instead of .get() for sqlite3.Row
            rca_data = json.loads(existing_rca["rca_json"])
            active_sessions[existing_rca["id"]] = {
                "rca_data": rca_data,
                "content": existing_rca["whiteboard_text"]
            }
            
            return {
                "rca_id": existing_rca["id"],
                "status": "CONFIRMED",
                "data": rca_data,
                "cached": True,
                "message": f"Retrieved confirmed RCA for {incident_id} from database"
            }
        elif existing_rca:
            print(f"   Found existing DRAFT for {incident_id}, regenerating...")
    
    # ============================================================
    # GENERATE NEW RCA
    # ============================================================
    print(f"🤖 Generating NEW RCA for {incident_id}...")
    rca_data = agent.process_incident(content, is_draft=True)
    rca_data["incident_id"] = incident_id
    
    whiteboard_content = rca_data.pop("_content", "") if "_content" in rca_data else content
    
    # CRITICAL FIX: Pass confirmed flag to save with correct status
    rca_id = store.save_incident(
        incident_id=incident_id,
        incident_type=incident_type,
        rca_data=rca_data,
        whiteboard_text=whiteboard_content,
        mir_text="",
        confirmed=confirm_save  # True = CONFIRMED, False = DRAFT
    )
    
    active_sessions[rca_id] = {
        "rca_data": rca_data,
        "content": content,
        "incident_id": incident_id
    }
    
    status_msg = "Confirmed & Saved" if confirm_save else "Draft (Preview only)"
    print(f"✅ {status_msg}: {rca_id}")
    
    return {
        "rca_id": rca_id,
        "status": "CONFIRMED" if confirm_save else "DRAFT",
        "data": rca_data,
        "cached": False,
        "confirm_required": not confirm_save,  # Tell frontend to show Save button
        "message": status_msg
    }

# ============================================================================
# ADD MIR (FINALIZE)
# ============================================================================

@app.post("/add-mir/{rca_id}")
async def add_mir(rca_id: str, mir_file: UploadFile = File(...)):
    print(f"\n📄 ADDING MIR to {rca_id}")
    print(f"   File: {mir_file.filename}")
    
    allowed_ext = ['.xlsx', '.xls', '.pdf', '.csv', '.docx']
    if not any(mir_file.filename.lower().endswith(ext) for ext in allowed_ext):
        raise HTTPException(400, f"Invalid file type. Allowed: {', '.join(allowed_ext)}")
    
    try:
        mir_bytes = await mir_file.read()
        if len(mir_bytes) > MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(400, f"MIR file too large")
        
        loop = asyncio.get_event_loop()
        mir_text = await loop.run_in_executor(
            None, 
            lambda: process_file(mir_bytes, mir_file.filename)
        )
        
        if not mir_text or len(mir_text.strip()) < 50:
            raise HTTPException(400, "Could not extract text from MIR file")
        
        print(f"   ✓ MIR processed: {len(mir_text)} chars")
        
        row = store.get(rca_id)
        if not row:
            raise HTTPException(404, "RCA not found")
        
        draft_data = json.loads(row["rca_json"])
        whiteboard_content = row["whiteboard_text"]  # FIXED: Was .get()
        
        final_rca = agent.process_incident(whiteboard_content, mir_text, is_draft=False)
        store.finalize(rca_id, {
        "final_root_cause": final_rca.get("probable_root_cause"),
        "corrective_actions": final_rca["sections"].get("corrective_actions", []),
         "preventive_actions": final_rca["sections"].get("preventive_actions", []),
        "finalized_at": datetime.now().isoformat()
        })
        final_rca["incident_id"] = draft_data["incident_id"]
        
        store.update_with_mir(rca_id, final_rca, mir_text)
        
        if rca_id in active_sessions:
            active_sessions[rca_id]["rca_data"] = final_rca
            active_sessions[rca_id]["mir"] = mir_text
        
        print(f"✅ Finalized with MIR")
        
        return {
            "rca_id": rca_id,
            "status": "FINAL",
            "data": final_rca
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Error: {e}")
        raise HTTPException(500, f"MIR processing failed: {str(e)}")

# ============================================================================
# OTHER ENDPOINTS
# ============================================================================

@app.post("/update-rca/{rca_id}")
async def update_rca(rca_id: str, data: dict):
    rca_data = data.get("rca_data")
    if not rca_data:
        raise HTTPException(400, "No RCA data provided")
    
    store.update_rca(rca_id, rca_data)
    
    if rca_id in active_sessions:
        active_sessions[rca_id]["rca_data"] = rca_data
    
    return {"success": True, "message": "RCA updated"}

@app.post("/finalize/{rca_id}")
async def finalize_rca(rca_id: str, data: dict):
    final_root = data.get("final_root_cause", "").strip()
    corrective = data.get("corrective_actions", [])
    preventive = data.get("preventive_actions", [])
    
    if not final_root:
        raise HTTPException(400, "Final root cause required")
    
    store.finalize(rca_id, {
        "final_root_cause": final_root,
        "corrective_actions": corrective,
        "preventive_actions": preventive,
        "finalized_at": datetime.now().isoformat()
    })
    
    return {"success": True, "rca_id": rca_id}

@app.post("/chat/{rca_id}")
async def chat(rca_id: str, data: dict):
    message = data.get("message", "").strip()
    
    if rca_id not in active_sessions:
        row = store.get(rca_id)
        if not row:
            raise HTTPException(404, "Not found")
        active_sessions[rca_id] = {
            "rca_data": json.loads(row["rca_json"]),
            "content": row["whiteboard_text"]  # FIXED: Was .get()
        }
    
    rca_data = active_sessions[rca_id]["rca_data"]
    
    msg_lower = message.lower()
    if "heading" in msg_lower:
        return {"response": f"Heading: {rca_data.get('heading')}"}
    elif "team" in msg_lower:
        teams = rca_data.get("incident_details", {}).get("teams", [])
        return {"response": f"Teams: {', '.join(teams)}"}
    
    return {"response": "I can help refine the RCA. What to modify?"}

@app.get("/rca/{rca_id}")
async def get_rca(rca_id: str):
    row = store.get(rca_id)
    if not row:
        raise HTTPException(404, "Not found")
    
    return {
        "rca_id": row["id"],  # FIXED: Was .get()
        "incident_id": row["incident_id"],
        "status": row["status"],
        "data": json.loads(row["rca_json"]) if row["rca_json"] else {},
        "has_mir": bool(row["mir_text"]),  # FIXED: Was .get()
        "created_at": row["created_at"]
    }

@app.get("/rca-history")
async def get_history():
    rows = store.list_all()
    result = []
    for r in rows:
        heading = ""
        if r["rca_json"]:  # Check if not None
            try:
                heading = json.loads(r["rca_json"]).get("heading", "")
            except:
                pass
        
        result.append({
            "id": r["id"],
            "incident_id": r["incident_id"],
            "heading": heading,
            "status": r["status"],
            "has_mir": bool(r["mir_text"]),  # FIXED
            "created_at": r["created_at"]
        })
    return result

if __name__ == "__main__":
    import uvicorn
    print("🚀 RCA Engine v5.1 (With Smart Caching)")
    print("📖 http://localhost:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)